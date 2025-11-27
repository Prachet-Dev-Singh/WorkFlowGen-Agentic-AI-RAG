from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.core.db import get_db
from app.models.document import Document, DocumentChunk
from app.schemas import DocumentResponse, QARequest, QAResponse, SummaryRequest, SummaryResponse
from app.services.llm import get_embedding, get_answer, get_summary
from app.agents.graph import app as agent_app

# New imports for file handling
import io
from pypdf import PdfReader
from docx import Document as DocxDocument

router = APIRouter()

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(file: UploadFile = File(...), db: AsyncSession = Depends(get_db)):
    filename = file.filename.lower()
    content = await file.read()
    text_content = ""

    try:
        # 1. Handle PDF Files
        if filename.endswith(".pdf"):
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            # Extract text from all pages
            text_content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            
        # 2. Handle Word Documents (.docx)
        elif filename.endswith(".docx"):
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            # Extract text from paragraphs
            text_content = "\n".join([para.text for para in doc.paragraphs])
            
        # 3. Handle Plain Text Files (.txt, .md, .csv)
        else:
            text_content = content.decode("utf-8")

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing file: {str(e)}")

    # Check if we actually extracted text
    if not text_content.strip():
         raise HTTPException(status_code=400, detail="The file appears to be empty or text could not be extracted (scanned PDFs are not supported yet).")

    # 4. Save to Database (Parent Record)
    new_doc = Document(title=file.filename, content=text_content)
    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)

    # 5. Chunking Logic (Split by 1000 characters for embedding)
    chunk_size = 1000
    chunks = [text_content[i:i+chunk_size] for i in range(0, len(text_content), chunk_size)]

    # 6. Create Vector Embeddings
    for index, chunk_text in enumerate(chunks):
        if not chunk_text.strip():
            continue
            
        vector = await get_embedding(chunk_text)
        
        new_chunk = DocumentChunk(
            document_id=new_doc.id,
            chunk_index=index,
            chunk_text=chunk_text,
            embedding=vector
        )
        db.add(new_chunk)
    
    await db.commit()
    await db.refresh(new_doc) # Refresh to ensure ID is available
    
    return new_doc

@router.get("/documents", response_model=list[DocumentResponse])
async def list_documents(db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document))
    return result.scalars().all()

# --- AGENT ENDPOINTS ---

@router.post("/agent", response_model=QAResponse)
async def run_agent(request: QARequest):
    # Runs the LangGraph Workflow
    result = await agent_app.ainvoke({"question": request.question})
    return QAResponse(
        answer=result["answer"],
        sources=["Agentic Search"]
    )

@router.post("/summarize", response_model=SummaryResponse)
async def summarize_document(request: SummaryRequest, db: AsyncSession = Depends(get_db)):
    stmt = select(DocumentChunk).where(DocumentChunk.document_id == request.document_id).order_by(DocumentChunk.chunk_index)
    result = await db.execute(stmt)
    chunks = result.scalars().all()
    
    if not chunks:
            raise HTTPException(status_code=404, detail="Document not found")

    full_text = "\n".join([c.chunk_text for c in chunks])
    
    # Generate Summary via AI
    summary = await get_summary(full_text)
    
    return SummaryResponse(summary=summary)

# Keep the old QA endpoint for backward compatibility (optional)
@router.post("/qa", response_model=QAResponse)
async def ask_question(request: QARequest, db: AsyncSession = Depends(get_db)):
    # ... (You can leave this or remove it, the Agent endpoint replaces it) ...
    # For now, let's just reuse the Agent logic here to keep it simple
    result = await agent_app.ainvoke({"question": request.question})
    return QAResponse(answer=result["answer"], sources=[])